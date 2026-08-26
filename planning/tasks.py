import io
import uuid
from datetime import datetime

from celery import shared_task
from django.conf import settings
from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.core.mail import EmailMessage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor

from planning.models import Activite, PlanAction, QuarterlyReport


@shared_task
def check_activity_alerts_task():
    Activite.check_activity_alerts()


User = get_user_model()


def _set_run_font(run, name='Arial', size=11, bold=False, italic=False, color=None):
    font = run.font
    font.name = name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    if color:
        font.color.rgb = color


def _set_cell_shading(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    )


CIV_ORANGE = RGBColor(0xFF, 0x9A, 0x00)
CIV_GREEN = RGBColor(0x00, 0x96, 0x39)
CIV_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CIV_BLACK = RGBColor(0x00, 0x00, 0x00)


def _build_report_doc(plan, annee, trimestre, activites_trimestre):
    annee_index = annee - plan.annee_debut
    today = datetime.now().strftime('%d/%m/%Y')

    doc = Document()

    # Police par défaut
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    # Justification par défaut
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # En-tête sobre
    header = section.header
    h1 = header.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(h1.add_run('RÉPUBLIQUE DE CÔTE D\'IVOIRE'), size=9, bold=True)
    h2 = header.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(h2.add_run('Union - Discipline - Travail'), size=8, italic=True)
    h3 = header.add_paragraph()
    h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(h3.add_run('Ministère des Finances et du Budget - DPSE MFB'), size=8)

    # Pied de page
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        footer_para.add_run(f'Document confidentiel - Généré le {today}'),
        size=8,
        color=RGBColor(0x66, 0x66, 0x66)
    )

    # Page de garde
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run('RÉPUBLIQUE DE CÔTE D\'IVOIRE'), size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run('Union - Discipline - Travail'), size=11, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run('MINISTÈRE DES FINANCES ET DU BUDGET'), size=13, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run('Direction du Pilotage et du Suivi-Évaluation'), size=12, bold=True)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run(f'RAPPORT TRIMESTRIEL {trimestre} {annee}'), size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run(f'Plan d\'actions : {plan.titre} ({plan.reference})'), size=12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run(f'Année {annee} - {trimestre}'), size=12)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(p.add_run(f'Abidjan, le {today}'), size=11, italic=True)

    # Section I. Introduction
    h = doc.add_heading('I. INTRODUCTION', level=1)
    _set_run_font(h.runs[0], size=14, bold=True, color=CIV_GREEN)

    p = doc.add_paragraph()
    p.add_run(
        f'Le présent rapport a été élaboré dans le cadre du dispositif de suivi-évaluation des plans d\'actions opérationnels du Ministère des Finances et du Budget. Il porte sur le {trimestre} de l\'année {annee} du plan d\'actions {plan.reference} intitulé « {plan.titre} ». Son objectif est de rendre compte, de manière factuelle et circonstanciée, de l\'état d\'avancement des activités programmées au cours de la période, d\'apprécier leur coût de mise en œuvre et de formuler, le cas échéant, des observations utiles à la poursuite des travaux. Ce document s\'adresse à l\'ensemble des parties prenantes impliquées dans la mise en œuvre, le suivi et la validation desdites activités.'
    )

    # Section II. Contexte
    h = doc.add_heading('II. CONTEXTE DE LA PÉRIODE', level=1)
    _set_run_font(h.runs[0], size=14, bold=True, color=CIV_GREEN)

    p = doc.add_paragraph()
    p.add_run(
        f'Le plan d\'actions {plan.reference} couvre un horizon de {plan.horizon} ans, à compter de l\'année {plan.annee_debut}. Pour l\'année {annee}, le {trimestre} constitue une période déterminante pour le lancement et le suivi des activités prioritaires. Les dispositions retenues pour la période visent à concourir aux effets escomptés du plan, notamment {plan.impact or "l\'amélioration de la gestion des finances publiques"}.'
    )

    # Section III. Etat d'avancement
    h = doc.add_heading('III. ÉTAT D\'AVANCEMENT DES ACTIVITÉS', level=1)
    _set_run_font(h.runs[0], size=14, bold=True, color=CIV_GREEN)

    total = len(activites_trimestre)
    par_statut = {}
    for a in activites_trimestre:
        s = a.status[annee_index] if a.status and len(a.status) > annee_index else 'Non entamée'
        par_statut[s] = par_statut.get(s, 0) + 1

    p = doc.add_paragraph()
    if total == 0:
        p.add_run(f'Pour le {trimestre} {annee}, aucune activité n\'a été programmée dans le plan d\'actions {plan.reference}.')
    else:
        p.add_run(f'Pour le {trimestre} {annee}, un total de {total} activité{"s" if total > 1 else ""} {"ont été" if total > 1 else "a été"} programmée{"s" if total > 1 else ""} au titre du plan d\'actions {plan.reference}. L\'analyse du niveau d\'avancement fait ressortir les éléments suivants :')
        for s, c in par_statut.items():
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'{c} activité{"s" if c > 1 else ""} au statut « {s} » ;')

    p = doc.add_paragraph()
    p.add_run('L\'examen des activités conduites durant la période permet de relever les éléments figurant dans le tableau ci-après :')

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    cols = ['Référence', 'Titre', 'Structure', 'Cible', 'Réalisation', 'Statut', 'Coût']
    for i, text in enumerate(cols):
        hdr[i].text = text
        for r in hdr[i].paragraphs[0].runs:
            _set_run_font(r, bold=True)
    for a in activites_trimestre:
        e = (a.point_focal.entity if a.point_focal and getattr(a.point_focal, 'entity', None) else
             a.responsable.entity if a.responsable and getattr(a.responsable, 'entity', None) else 'Sans entité')
        row = table.add_row().cells
        row[0].text = a.reference or ''
        row[1].text = a.titre or ''
        row[2].text = e
        row[3].text = a.cibles[annee_index] if a.cibles and len(a.cibles) > annee_index else ''
        row[4].text = a.realisation[annee_index] if a.realisation and len(a.realisation) > annee_index else ''
        row[5].text = a.status[annee_index] if a.status and len(a.status) > annee_index else 'Non entamée'
        row[6].text = f"{float(a.couts[annee_index] or 0):,.2f}" if a.couts and len(a.couts) > annee_index else '0.00'

    if total > 0:
        p = doc.add_paragraph()
        p.add_run('En termes de conduite, il convient de souligner les points suivants :')
        for a in activites_trimestre:
            e = (a.point_focal.entity if a.point_focal and getattr(a.point_focal, 'entity', None) else
                 a.responsable.entity if a.responsable and getattr(a.responsable, 'entity', None) else 'Sans entité')
            st = a.status[annee_index] if a.status and len(a.status) > annee_index else 'Non entamée'
            cible = a.cibles[annee_index] if a.cibles and len(a.cibles) > annee_index else ''
            real = a.realisation[annee_index] if a.realisation and len(a.realisation) > annee_index else ''
            cout = float(a.couts[annee_index] or 0) if a.couts and len(a.couts) > annee_index else 0.0
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'L\'activité {a.reference or a.titre}, conduite par la structure {e}, poursuit l\'objectif suivant : {a.titre}. La cible fixée était « {cible} » et le niveau de réalisation enregistré s\'élève à « {real} ». Le statut actuel est « {st} », pour un coût de mise en œuvre de {cout:,.2f}.')

    # Section IV. Synthèse financière
    h = doc.add_heading('IV. SYNTHÈSE FINANCIÈRE', level=1)
    _set_run_font(h.runs[0], size=14, bold=True, color=CIV_GREEN)

    total_cout = sum(float(a.couts[annee_index] or 0) for a in activites_trimestre if a.couts and len(a.couts) > annee_index)
    p = doc.add_paragraph()
    if total_cout:
        p.add_run(f'La mise en œuvre des activités du {trimestre} {annee} a nécessité un montant total de {total_cout:,.2f}. Ce montant se décompose par structure de la manière suivante :')
    else:
        p.add_run(f'Pour le {trimestre} {annee}, aucun coût n\'a été enregistré pour les activités retenues.')

    by_entity = {}
    for a in activites_trimestre:
        e = (a.point_focal.entity if a.point_focal and getattr(a.point_focal, 'entity', None) else
             a.responsable.entity if a.responsable and getattr(a.responsable, 'entity', None) else 'Sans entité')
        c = by_entity.setdefault(e, {'cout': 0.0, 'count': 0})
        c['count'] += 1
        c['cout'] += float(a.couts[annee_index] or 0) if a.couts and len(a.couts) > annee_index else 0.0

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Structure'
    hdr[1].text = 'Coût'
    hdr[2].text = 'Activités'
    for cell in hdr:
        for r in cell.paragraphs[0].runs:
            _set_run_font(r, bold=True)
    for e, c in sorted(by_entity.items()):
        row = table.add_row().cells
        row[0].text = e
        row[1].text = f"{c['cout']:,.2f}"
        row[2].text = str(c['count'])

    p = doc.add_paragraph()
    p.add_run('Cette répartition fait apparaître la répartition des efforts financiers entre les structures en charge de la mise en œuvre. Elle constitue un indicateur utile pour apprécier la concentration des moyens et l\'équité de l\'allocation des ressources.')

    # Section V. Constats et recommandations
    h = doc.add_heading('V. CONSTATS ET RECOMMANDATIONS', level=1)
    _set_run_font(h.runs[0], size=14, bold=True, color=CIV_GREEN)

    p = doc.add_paragraph()
    if total == 0:
        p.add_run(f'Pour le {trimestre} {annee}, aucune activité n\'ayant été programmée dans le plan d\'actions {plan.reference}, il est recommandé de veiller à la planification effective des activités des prochains trimestres afin de garantir la cohérence et la régularité de la mise en œuvre.')
    else:
        p.add_run(f'L\'analyse des activités du {trimestre} {annee} fait ressortir un ensemble de tendances. Le nombre d\'activités en cours ou réalisées indique un niveau d\'engagement des structures. Il est recommandé de poursuivre le suivi régulier des activités programmées, de mobiliser les ressources nécessaires à leur achèvement et de rapporter, en temps utile, les difficultés rencontrées. Le respect des échéances trimestrielles et la qualité des données de suivi demeurent des facteurs essentiels pour l\'efficacité du dispositif de pilotage.')

    p = doc.add_paragraph()
    p.add_run('Afin de renforcer la performance du plan d\'actions, il est proposé :')
    for rec in [
        'd\'accorder une attention particulière aux activités en retard et de mettre en œuvre les mesures correctives nécessaires ;',
        'de veiller à la régularité de la saisie des preuves de réalisation et des éléments justificatifs ;',
        'de tenir les réunions de revue trimestrielle prévues pour l\'examen des progrès et l\'ajustement des actions si nécessaire ;',
        'de diffuser, en interne, les présents résultats aux responsables et points focaux concernés.'
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(rec)

    # Conclusion
    h = doc.add_heading('VI. CONCLUSION', level=1)
    _set_run_font(h.runs[0], size=14, bold=True, color=CIV_GREEN)

    p = doc.add_paragraph()
    if total == 0:
        p.add_run(f'En conclusion, le {trimestre} {annee} n\'a enregistré aucune activité au titre du plan d\'actions {plan.reference}. Les prochains trimestres offriront l\'opportunité de rattraper le retard et de déployer effectivement les actions retenues.')
    else:
        p.add_run(f'En conclusion, le {trimestre} {annee} a vu la mise en œuvre de {total} activité{"s" if total > 1 else ""} au titre du plan d\'actions {plan.reference}. Les résultats enregistrés, bien qu\'encourageants, appellent à une vigilance soutenue afin de garantir la finalisation des actions en cours et la réalisation des objectifs annuels. Le présent rapport, qui sera régulièrement actualisé, constitue un outil de travail pour le pilotage et l\'orientation des décisions à venir.')

    return doc


@shared_task
def generate_quarterly_report(plan_id, annee, trimestre, user_id=None):
    plan = PlanAction.objects.get(id=plan_id)
    annee_index = annee - plan.annee_debut
    if annee_index < 0 or annee_index >= plan.horizon:
        raise ValueError("Année hors de l'horizon du plan.")

    activites = Activite.objects.filter(
        action__produit__effet__plan=plan
    ).select_related('action__produit__effet', 'point_focal', 'responsable')

    activites_trimestre = []
    for a in activites:
        if not a.periodes_execution or len(a.periodes_execution) <= annee_index:
            continue
        if trimestre in a.periodes_execution[annee_index]:
            activites_trimestre.append(a)

    doc = _build_report_doc(plan, annee, trimestre, activites_trimestre)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    content = buffer.read()

    filename = f"rapport_{plan.reference}_{annee}_{trimestre}_{uuid.uuid4().hex[:8]}.docx"

    report, _ = QuarterlyReport.objects.get_or_create(
        plan=plan,
        annee=annee,
        trimestre=trimestre,
        defaults={'generated_by_id': user_id, 'recipients_emails': []}
    )

    if report.fichier:
        report.fichier.delete(save=False)

    report.fichier.save(filename, ContentFile(content), save=True)
    report.generated_by_id = user_id

    # Destinataires : acteurs du plan
    recipients = set()
    for a in activites_trimestre:
        if a.point_focal and a.point_focal.email:
            recipients.add(a.point_focal.email)
        if a.responsable and a.responsable.email:
            recipients.add(a.responsable.email)
    for u in User.objects.filter(is_active=True, groups__name='SuiveurEvaluateur'):
        if u.email:
            recipients.add(u.email)

    report.recipients_emails = sorted(recipients)
    report.save(update_fields=['recipients_emails', 'generated_by'])

    if recipients:
        email = EmailMessage(
            subject=f"Rapport trimestriel {trimestre} {annee} - {plan.reference}",
            body=f"Veuillez trouver ci-joint le rapport trimestriel {trimestre} {annee} du plan {plan.reference}.",
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=list(recipients),
        )
        email.attach_file(report.fichier.path)
        email.send(fail_silently=True)

    return report.id
